import os
import random
import docx
from pathlib import Path

from app.utils.utils import *
from app.models import Document
from app import db
import app.routes as routes

random.seed(420)


def delete_unused_records():
    documents = Document.query.all()
    unused_documents = [document for document in documents if not os.path.exists(document.abs_path)]

    if unused_documents:
        unused_ids = [doc.id for doc in unused_documents]
        db.session.query(Document).filter(Document.id.in_(unused_ids)).delete(synchronize_session=False)
        db.session.commit()


def extract_text_from_docx(path):
    try:
        doc = docx.Document(path)
        return "\n".join(para.text for para in doc.paragraphs)
    except Exception as e:
        print(f"Error reading docx {path}: {e}")
        return ""


def extract_text(path, file_type):
    try:
        if file_type in ["pdf", "jpeg", "jpg", "png"]:
            return img_to_text(path)
        elif file_type in ["doc", "docx"]:
            return extract_text_from_docx(path)
        else:
            with open(path, "r", encoding="utf-8") as file:
                return file.read().strip()
    except Exception as e:
        print(f"Error extracting text from {path}: {e}")
        return ""


def get_tags(file_type, abs_path):
    if file_type != "folder":
        text = extract_text(abs_path, file_type)
        return str([open_ai_model(text)])
    return "['None']"


def get_icons(file_type):
    icon_class = f"bi bi-filetype-{file_type}" if file_type in file_type_list else "bi bi-folder"
    return icon_class


def get_obj_from_scan(root, name):
    file_type = Path(name).suffix
    file_type = file_type[1:] if file_type.startswith(".") else "folder"

    abs_path = os.path.join(root, name).replace("\\", "/")
    rel_path = os.path.relpath(abs_path, routes.chosen_folder_path).replace("\\", "/")

    stat = os.stat(abs_path)
    inode = stat.st_ino
    bytes = get_readable_byte_size(stat.st_size)
    m_time = get_time_stamp_string(stat.st_mtime)
    icon = get_icons(file_type)

    return Document(
        id=inode,
        name=name,
        type=file_type,
        size=bytes,
        abs_path=abs_path,
        rel_path=rel_path,
        date_modified=m_time,
        icon=icon,
        tags="['None']",
    )


def is_modified_document(document, document_obj):
    if document.size != document_obj.size or document.date_modified != document_obj.date_modified:
        return 0
    elif document.abs_path != document_obj.abs_path or document.rel_path != document_obj.rel_path:
        return 1
    return 2


def add_or_update_docment(documents_to_add, documents_to_update, root, name):
    document_obj = get_obj_from_scan(
        root,
        name,
    )

    document = Document.query.filter_by(id=document_obj.id).first()
    if document:
        if is_modified_document(document, document_obj) == 0:
            document_obj.tags = get_tags(document_obj.type, document_obj.abs_path)
            documents_to_update.append(document_obj)
        elif is_modified_document(document, document_obj) == 1:
            document_obj.tags = document.tags
            documents_to_update.append(document_obj)
    else:
        document_obj.tags = get_tags(document_obj.type, document_obj.abs_path)
        documents_to_add.append(document_obj)


def upload_documents_to_db():
    documents_to_add = []
    documents_to_update = []

    for root, dirs, files in os.walk(routes.chosen_folder_path):
        for dir_name in dirs:
            add_or_update_docment(documents_to_add, documents_to_update, root, dir_name)
        for file_name in files:
            add_or_update_docment(documents_to_add, documents_to_update, root, file_name)

    if documents_to_add:
        db.session.bulk_save_objects(documents_to_add)
    if documents_to_update:
        db.session.bulk_update_mappings(Document, [doc.__dict__ for doc in documents_to_update])
    db.session.commit()
